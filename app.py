import streamlit as st
import google.generativeai as genai
import requests
from bs4 import BeautifulSoup
import re
import datetime
import concurrent.futures
import os
import markdown
from docx import Document
from docx.shared import Pt, RGBColor
from io import BytesIO

# ==========================================
# 0. 【网络配置】
# ==========================================
if "OS" in os.environ:
    os.environ["http_proxy"] = "http://127.0.0.1:7897"
    os.environ["https_proxy"] = "http://127.0.0.1:7897"

# ==========================================
# 1. 配置区域
# ==========================================
try:
    GEMINI_API_KEY = st.secrets["GEMINI_API_KEY"]
except FileNotFoundError:
    st.error("❌ 未找到密钥配置！请确保本地有 .streamlit/secrets.toml 或云端已配置 Secrets。")
    st.stop()

genai.configure(api_key=GEMINI_API_KEY)
model = genai.GenerativeModel('gemini-flash-latest')

# ==========================================
# 2. 核心功能
# ==========================================
def extract_urls(text):
    url_pattern = r'(https?://[^\s]+)'
    return re.findall(url_pattern, text)

def scrape_one_url(url):
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
    }
    try:
        response = requests.get(url, headers=headers, timeout=10)
        if response.status_code == 200:
            soup = BeautifulSoup(response.text, 'html.parser')
            paragraphs = soup.find_all('p')
            text = "\n".join([p.get_text() for p in paragraphs])
            if len(text) < 50: text = soup.get_text()
            return url, text[:2500] 
        return url, f"[网页错误: {response.status_code}]"
    except Exception as e:
        return url, f"[抓取出错: {str(e)}]"

def ai_generate_daily_brief(raw_input, scraped_text_block, lang_mode):
    
    # 基础结构：要求 AI 严格分行
    base_prompt = f"""
    你是一位 Chilquinta 能源公司的情报专家。
    请根据提供的【原始分类】和【抓取的正文】，写一份排版精美的日报。
    时间：{datetime.date.today()}
    
    【排版严格要求】：
    每一条新闻必须严格遵守以下 Markdown 结构（注意空行）：
    
    ### 🍊 [标题]
    
    [正文段落1]
    
    [正文段落2 (如果是双语模式)]
    
    **🔗 Source:** [URL]
    
    ---
    """

    if lang_mode == "中文 (保留西语术语)":
        lang_instruction = """
        【语言要求】：
        1. 标题：中文。
        2. 正文：中文摘要。
        3. **术语保留**：机构名、法规、项目名、专有名词后必须保留西语原文，如：国家能源委员会 (CNE)。
        """
    elif lang_mode == "纯西语 (Español)":
        lang_instruction = """
        【语言要求】：
        1. 标题：Español.
        2. 正文：Resumen en Español (Formal Business Tone).
        """
    else: # 中文 & 西语对照 (🌟 修改点：这里加上了术语保留的要求)
        lang_instruction = """
        【语言要求 - 双语对照模式】：
        请严格按以下格式输出，不要把中西文混在一段里：
        
        **🇨🇳 中文摘要：**
        [这里写中文摘要。⚠️关键要求：机构名、法规、项目名、专有名词后必须保留西语原文，例如：国家能源委员会 (CNE)、行政部门 (Ejecutivo)。]
        
        **🇪🇸 Español:**
        [Aquí el resumen en español...]
        """

    full_prompt = base_prompt + lang_instruction
    
    try:
        full_content = f"【原始消息框架】:\n{raw_input}\n\n【抓取的详细正文】:\n{scraped_text_block}"
        response = model.generate_content(full_prompt + "\n\n" + full_content)
        return response.text
    except Exception as e:
        return f"AI 思考出错: {str(e)}"

def convert_to_html_file(markdown_text):
    html_body = markdown.markdown(markdown_text)
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <style>
            body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Helvetica, Arial, sans-serif; line-height: 1.6; color: #333; background-color: #f4f7f6; max-width: 800px; margin: 0 auto; padding: 20px; }}
            .container {{ background-color: #ffffff; padding: 40px; border-radius: 12px; box-shadow: 0 4px 15px rgba(0,0,0,0.05); }}
            h2 {{ color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; margin-top: 40px; margin-bottom: 25px; font-size: 1.5em; font-weight: 700; }}
            h3 {{ color: #d35400; margin-top: 30px; margin-bottom: 15px; font-size: 1.2em; font-weight: 600; }}
            p {{ margin-bottom: 15px; font-size: 15px; color: #444; text-align: justify; }}
            strong {{ color: #2c3e50; font-weight: 600; }}
            ul {{ background-color: #f8f9fa; padding: 10px 15px 10px 35px; border-radius: 6px; border-left: 4px solid #3498db; margin-bottom: 25px; }}
            li {{ margin-bottom: 5px; font-size: 13px; color: #666; word-break: break-all; }}
            a {{ color: #007bff; text-decoration: none; }}
            .footer {{ margin-top: 40px; text-align: center; font-size: 12px; color: #aaa; }}
            @media only screen and (max-width: 600px) {{
                body {{ padding: 10px; }}
                .container {{ padding: 20px; }}
                h2 {{ font-size: 1.3em; }}
                h3 {{ font-size: 1.1em; }}
                p {{ font-size: 15px; }}
            }}
        </style>
    </head>
    <body>
        <div class="container">{html_body}<div class="footer">⚡ Generated by Chilquinta AI Assistant • {datetime.date.today()}</div></div>
    </body>
    </html>
    """
    return html_content

def generate_word_file(markdown_text):
    doc = Document()
    doc.add_heading(f'Chilquinta Daily News - {datetime.date.today()}', 0)
    lines = markdown_text.split('\n')
    for line in lines:
        line = line.strip()
        if not line: continue
        if line.startswith('### '):
            clean_line = line.replace('### ', '').replace('🍊', '').strip()
            heading = doc.add_heading(clean_line, level=2)
            run = heading.runs[0]
            run.font.color.rgb = RGBColor(211, 84, 0)
        elif line.startswith('* ') or line.startswith('- '):
            clean_line = re.sub(r'^[*-]\s+', '', line)
            clean_line = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', clean_line) 
            doc.add_paragraph(clean_line, style='List Bullet')
        elif "Source:" in line or "🔗" in line:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(100, 100, 100)
        elif line.startswith('---'):
            doc.add_paragraph('_' * 20)
        else:
            clean_line = line.replace('**', '')
            doc.add_paragraph(clean_line)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==========================================
# 3. 界面构建
# ==========================================
st.set_page_config(page_title="Chilquinta News v1.2", page_icon="⚡", layout="wide")

# 🌟 新增：侧边栏更新日志
with st.sidebar:
    st.header("📅 更新日志")
    st.markdown("""
    * **2026.01.15 (v1.2 修补)**
        * 优化双语模式：中文部分现已包含西语术语对照。
    * **2026.01.13 (v1.2)**
        * 优化排版：强制分行，字体层级优化（护眼模式）。
    * **2026.01.13 (v1.1)**
        * 新增功能：Word下载、多语言切换。
    * **2025.12.24 (v1.0)**
        * 初始版本发布。
    """)
    st.info("💡 提示：双语模式下，中文摘要也会保留关键术语（如 Ejecutivo），方便对照。")

st.title("⚡ Chilquinta 每日新闻 (v1.2)")
st.caption("排版优化版 • 字体层级清晰 • 独立分行")

raw_text = st.text_area("请粘贴新闻链接:", height=150)

lang_option = st.radio(
    "请选择生成语言:",
    ("中文 (保留西语术语)", "纯西语 (Español)", "中文 & 西语对照"),
    horizontal=True
)

if st.button("🚀 开始生成日报", type="primary"):
    if not raw_text or "http" not in raw_text:
        st.warning("请粘贴包含链接的内容！")
    else:
        urls = extract_urls(raw_text)
        status = st.status(f"发现 {len(urls)} 条链接，正在并发抓取...", expanded=True)
        scraped_data_str = ""
        
        with concurrent.futures.ThreadPoolExecutor(max_workers=10) as executor:
            future_to_url = {executor.submit(scrape_one_url, url): url for url in urls}
            for future in concurrent.futures.as_completed(future_to_url):
                url = future_to_url[future]
                try:
                    _, content = future.result()
                    scraped_data_str += f"\n--- 链接 {url} 的正文 ---\n{content}\n"
                    status.write(f"✅ 已抓取: {url[:40]}...")
                except:
                    status.write(f"❌ 失败: {url[:40]}")
        
        status.write(f"🧠 AI 正在用【{lang_option}】模式撰写报告...")
        
        report_md = ai_generate_daily_brief(raw_text, scraped_data_str, lang_option)
        report_html = convert_to_html_file(report_md)
        word_file = generate_word_file(report_md)
        
        status.update(label="✅ 完成！", state="complete", expanded=False)
        st.markdown("---")
        
        col1, col2 = st.columns(2)
        with col1:
            st.download_button(
                label="📥 下载网页版 (.html)",
                data=report_html,
                file_name=f"Chilquinta_News_{datetime.date.today()}.html",
                mime="text/html"
            )
        with col2:
            st.download_button(
                label="📥 下载 Word 版 (.docx)",
                data=word_file,
                file_name=f"Chilquinta_News_{datetime.date.today()}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
        st.markdown(report_md)